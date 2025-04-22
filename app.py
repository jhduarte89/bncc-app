import streamlit as st
import docx
from io import BytesIO

competencias = {
    "CE1": {
        "descricao": "Analisar processos políticos, econômicos, sociais, ambientais e culturais nos âmbitos local, regional, nacional e mundial.",
        "habilidades": {
            "EM13CHS101": "Identificar, analisar e comparar diferentes fontes e narrativas expressas em diversas linguagens.",
            "EM13CHS102": "Analisar circunstâncias históricas de matrizes conceituais como etnocentrismo, racismo, etc.",
            "EM13CHS103": "Elaborar hipóteses e compor argumentos com base em dados e evidências.",
            "EM13CHS104": "Analisar objetos e vestígios da cultura material e imaterial.",
            "EM13CHS105": "Criticar oposições dicotômicas como cidade/campo, civilizados/bárbaros, etc.",
            "EM13CHS106": "Utilizar linguagens cartográfica e digital de forma crítica e ética."
        }
    }
}

st.set_page_config(page_title="BNCC: Competências e Habilidades", layout="centered")
st.title("🎓 Seletor de Competências e Habilidades - BNCC")
st.markdown("Escolha as competências e as habilidades que deseja incluir no seu documento.")

escolhas = {}
for cod, comp in competencias.items():
    with st.expander(f"{cod} - {comp['descricao']}"):
        sel = st.multiselect(
            f"Selecione as habilidades para {cod}:",
            options=list(comp["habilidades"].keys()),
            format_func=lambda x: f"{x} - {comp['habilidades'][x]}"
        )
        if sel:
            escolhas[cod] = sel

def gerar_documento(selecionadas):
    doc = docx.Document()
    doc.add_heading("Competências e Habilidades Selecionadas - BNCC", 0)
    for cod, habs in selecionadas.items():
        doc.add_heading(f"{cod} - {competencias[cod]['descricao']}", level=1)
        for hab in habs:
            doc.add_paragraph(f"{hab}: {competencias[cod]['habilidades'][hab]}", style='List Bullet')
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

if st.button("📄 Gerar Documento Word"):
    if escolhas:
        file = gerar_documento(escolhas)
        st.download_button("📥 Baixar Documento", data=file, file_name="competencias_habilidades_bncc.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.warning("Por favor, selecione pelo menos uma habilidade.")
